"""
Smart Report Engine v4.0 — Công cụ xuất báo cáo ĐATN chuẩn TDTU.
Phiên bản tổng quát: Hoạt động với MỌI đề tài.
Fix v4: Mermaid URL-safe encoding + retry, no bullet period, char spacing total fix.

Features:
- Trang bìa chính + bìa phụ chuẩn TDTU
- Lời cảm ơn + Lời cam đoan tự động
- TÓM TẮT / ABSTRACT
- Heading 3 cấp đúng quy định
- Công thức toán học có đánh số + giải thích biến
- Bảng biểu với tiêu đề phía trên
- Hình vẽ + chú thích phía dưới
- Mermaid diagram → ảnh PNG
- Chèn code snippet (Consolas 10pt)
- Cấu trúc thư mục dự án
- Đánh số trang (header, giữa)
- Lề chuẩn TDTU (3.5/3/3.5/2 cm)
"""

import os
import re
import base64
import requests
import io
import time
import urllib.parse
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


class ReportGenerator:
    """Bộ sinh báo cáo Word (.docx) chuẩn định dạng TDTU — Tổng quát mọi đề tài."""

    DEFAULTS = {
        "font_name": "Times New Roman",
        "font_size": 13,
        "heading1_size": 16,
        "heading2_size": 14,
        "heading3_size": 13,
        "code_font": "Consolas",
        "code_size": 10,
        "line_spacing": 1.5,
        "margin_top": Cm(3.5),
        "margin_bottom": Cm(3.0),
        "margin_left": Cm(3.5),
        "margin_right": Cm(2.0),
        "first_line_indent": Cm(1.27),
    }

    def __init__(self, title, author, advisor,
                 department="KHOA ĐIỆN – ĐIỆN TỬ",
                 major="KỸ THUẬT ĐIỆN TỬ - VIỄN THÔNG",
                 year="2026", mssv="", lop="", report_type="BÁO CÁO ĐỒ ÁN TỐT NGHIỆP NGÀNH KỸ THUẬT"):
        self.doc = Document()
        self.title = title
        self.author = author
        self.advisor = advisor
        self.department = department
        self.major = major
        self.year = year
        self.mssv = mssv
        self.lop = lop
        self.report_type = report_type
        self.current_chapter = 0
        self._setup_page()
        self._setup_styles()

    # ═══════════════ PAGE SETUP ═══════════════
    def _setup_page(self):
        """Thiết lập kích thước trang, lề cơ bản cho Section 0 (Trang bìa)."""
        section = self.doc.sections[0]
        self._apply_page_dimensions(section)
        
        # Ẩn header và footer ở Section 0
        section.different_first_page_header_footer = True

    def _apply_page_dimensions(self, section):
        from docx.enum.section import WD_ORIENT
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = self.DEFAULTS["margin_top"]
        section.bottom_margin = self.DEFAULTS["margin_bottom"]
        section.left_margin = self.DEFAULTS["margin_left"]
        section.right_margin = self.DEFAULTS["margin_right"]
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.5)

    def start_front_matter(self):
        """Bắt đầu phần Front Matter (Lời cảm ơn, Abstract...), sử dụng số La Mã ở giữa footer."""
        from docx.enum.section import WD_SECTION
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        self._apply_page_dimensions(new_section)
        new_section.is_linked_to_previous = False
        new_section.different_first_page_header_footer = False
        self._just_started_section = True # Prevent double page break for heading
        
        # Thiết lập số trang La Mã (i, ii, iii...)
        sectPr = new_section._sectPr
        pgNumType = parse_xml(r'<w:pgNumType xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fmt="lowerRoman" w:start="1"/>')
        sectPr.append(pgNumType)
        
        # Xóa Header
        header = new_section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.text = ""
            
        # Thêm Footer chứa số trang
        footer = new_section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run()
        run.font.name = self.DEFAULTS["font_name"]
        run.font.size = Pt(12)
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        run._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    def start_main_content(self):
        """Bắt đầu phần Nội dung chính, Header chứa tên đồ án và số trang, Footer chứa tên đề tài."""
        from docx.enum.section import WD_SECTION
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        self._apply_page_dimensions(new_section)
        new_section.is_linked_to_previous = False
        new_section.different_first_page_header_footer = False
        self._just_started_section = True # Prevent double page break for heading
        
        # Thiết lập số trang Decimal (1, 2, 3...)
        sectPr = new_section._sectPr
        pgNumType = parse_xml(r'<w:pgNumType xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fmt="decimal" w:start="1"/>')
        sectPr.append(pgNumType)
        
        # HEADER
        header = new_section.header
        header.is_linked_to_previous = False
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # ĐỒ ÁN TỐT NGHIỆP/ TỔNG HỢP \n Trang x
        r1 = header_para.add_run("ĐỒ ÁN TỐT NGHIỆP/ TỔNG HỢP\nTrang ")
        r1.font.name = self.DEFAULTS["font_name"]
        r1.font.size = Pt(12)
        r1.bold = True
        r1.font.color.rgb = RGBColor(100, 100, 100) # Mau xam theo chuan
        
        run = header_para.add_run()
        run.font.name = self.DEFAULTS["font_name"]
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(100, 100, 100)
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        run._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
        
        # Them Bottom border line cho Header
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>')
        header_para._p.get_or_add_pPr().append(pBdr)

        # FOOTER
        footer = new_section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Them Top border line cho Footer
        pBdr2 = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="6" w:space="1" w:color="auto"/></w:pBdr>')
        footer_para._p.get_or_add_pPr().append(pBdr2)
        
        f_run = footer_para.add_run(self.title if hasattr(self, 'title') else "")
        f_run.font.name = self.DEFAULTS["font_name"]
        f_run.font.size = Pt(12)
        f_run.italic = False
        f_run.bold = True
        f_run.font.color.rgb = RGBColor(100, 100, 100)

    def _setup_styles(self):
        """Thiết lập định dạng chuẩn cho toàn bộ văn bản."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.DEFAULTS["font_name"]
        font.size = Pt(self.DEFAULTS["font_size"])

        # Áp dụng cho cả font tiếng Việt
        rpr = style.element.xpath('.//w:rPr')
        if rpr:
            rpr[0].append(
                self.doc.element.makeelement(
                    qn('w:rFonts'),
                    {qn('w:eastAsia'): self.DEFAULTS["font_name"]}
                )
            )
            # Fix character spacing: đặt spacing = 0 để tránh dãn chữ
            spacing_elem = parse_xml(f'<w:spacing {nsdecls("w")} w:val="0"/>')
            rpr[0].append(spacing_elem)

        pf = style.paragraph_format
        pf.line_spacing = self.DEFAULTS["line_spacing"]
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = self.DEFAULTS["first_line_indent"]
        pf.space_after = Pt(6)  # Giãn dòng sau mỗi đoạn văn 6pt chuẩn TDTU

        # Fix: Tắt auto-spacing giữa ký tự Đông Á và Latin
        pPr = style.element.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:autoSpaceDE {nsdecls("w")} w:val="0"/>'))
        pPr.append(parse_xml(f'<w:autoSpaceDN {nsdecls("w")} w:val="0"/>'))
        pPr.append(parse_xml(f'<w:adjustRightInd {nsdecls("w")} w:val="0"/>'))

    def _fix_char_spacing(self, run):
        """Ép character spacing = 0 trên run để Word không dãn chữ khi Justify."""
        rPr = run._r.get_or_add_rPr()
        spacing = rPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = parse_xml(f'<w:spacing {nsdecls("w")} w:val="0"/>')
            rPr.append(spacing)
        else:
            spacing.set(qn('w:val'), '0')

    def _make_run(self, paragraph, text, bold=False, italic=False, size=None, font=None, color=None):
        """Helper: tạo Run với format chuẩn + fix character spacing."""
        run = paragraph.add_run(text)
        run.font.name = font or self.DEFAULTS["font_name"]
        run.font.size = Pt(size or self.DEFAULTS["font_size"])
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        self._fix_char_spacing(run)
        return run

    # ═══════════════ COVER PAGES ═══════════════
    def add_cover_page(self, include_advisor=False):
        """Trang bìa chuẩn TDTU với các tỉ lệ font chữ chính xác và Logo."""
        self.add_page_break()
        # Dòng 1: TỔNG LIÊN ĐOÀN...
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        self._make_run(p, "TỔNG LIÊN ĐOÀN LAO ĐỘNG VIỆT NAM", size=14)

        # Dòng 2: TDTU
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        self._make_run(p, "TRƯỜNG ĐẠI HỌC TÔN ĐỨC THẮNG", bold=True, size=14)

        # Dòng 3: Khoa
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        dept = self.department if hasattr(self, 'department') else "KHOA ĐIỆN - ĐIỆN TỬ"
        self._make_run(p, dept.upper(), bold=True, size=14)

        self.doc.add_paragraph()

        # Logo TDTU
        self.add_image_from_url("https://upload.wikimedia.org/wikipedia/vi/1/1b/TDTU_logo.png", width_inches=2.2)

        # Loại đồ án
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        report_type = self.report_type if hasattr(self, 'report_type') else "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP NGÀNH KỸ THUẬT"
        self._make_run(p, report_type.upper(), bold=True, size=18)

        self.doc.add_paragraph()

        # Tên đề tài
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        title = self.title if hasattr(self, 'title') else "TÊN ĐỀ TÀI ĐỒ ÁN"
        self._make_run(p, title.upper(), bold=True, size=22)

        for _ in range(2):
            self.doc.add_paragraph()

        # Người thực hiện
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        self._make_run(p, "Người thực hiện: ", size=14)
        author = self.author if hasattr(self, 'author') else "Tên sinh viên"
        self._make_run(p, author, bold=True, size=14)

        # Cố vấn
        if include_advisor:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            self._make_run(p, "Người hướng dẫn: ", size=14)
            advisor = self.advisor if hasattr(self, 'advisor') else "Tên giảng viên"
            self._make_run(p, advisor, bold=True, size=14)

        # Đẩy xuống cuối trang
        for _ in range(2):
            self.doc.add_paragraph()

        # Nơi xuất bản
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        year = self.year if hasattr(self, 'year') else "2026"
        self._make_run(p, f"THÀNH PHỐ HỒ CHÍ MINH, NĂM {year}", bold=True, size=14)

        

    # ═══════════════ FRONT MATTER ═══════════════
    def add_acknowledgment(self, text=None):
        """Trang Lời cảm ơn chuẩn TDTU."""
        self.add_page_break()
        self.add_heading("LỜI CẢM ƠN")
        if text is None:
            text = (
                f"Tác giả xin gửi lời cảm ơn chân thành và sâu sắc nhất đến {self.advisor} – "
                "người đã tận tình hướng dẫn, định hướng nghiên cứu và hỗ trợ "
                "trong suốt quá trình thực hiện đồ án này.\n\n"
                "Tác giả xin trân trọng cảm ơn quý thầy cô trong "
                f"{self.department} và Trường Đại học Tôn Đức Thắng "
                "đã truyền đạt những kiến thức quý báu trong suốt thời gian học tập tại trường.\n\n"
                "Cuối cùng, tác giả xin gửi lời cảm ơn đến gia đình, bạn bè "
                "đã luôn động viên, ủng hộ và tạo điều kiện tốt nhất để tác giả "
                "có thể hoàn thành đồ án tốt nghiệp này."
            )
        self.add_paragraph(text)

        # Chữ ký
        self.doc.add_paragraph()
        sig = self.doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig.paragraph_format.first_line_indent = None
        self._make_run(sig, f"TP. Hồ Chí Minh, ngày ... tháng ... năm {self.year}\n", italic=True)
        self._make_run(sig, "Tác giả\n\n\n", italic=True)
        self._make_run(sig, self.author, italic=True)
        

    def add_declaration(self):
        """Trang Lời cam đoan chuẩn TDTU."""
        self.add_page_break()
        self.add_heading("CÔNG TRÌNH ĐƯỢC HOÀN THÀNH TẠI TRƯỜNG ĐẠI HỌC TÔN ĐỨC THẮNG")
        text = (
            "Tôi xin cam đoan đây là công trình nghiên cứu của riêng tôi và được sự hướng dẫn "
            f"khoa học của {self.advisor}. "
            "Các nội dung nghiên cứu, kết quả trong đề tài này là trung thực "
            "và chưa công bố dưới bất kỳ hình thức nào trước đây. "
            "Những số liệu trong các bảng biểu phục vụ cho việc phân tích, nhận xét, "
            "đánh giá được chính tác giả thu thập từ các nguồn khác nhau có ghi rõ "
            "trong phần tài liệu tham khảo.\n\n"
            "Ngoài ra, trong đồ án còn sử dụng một số nhận xét, đánh giá cũng như "
            "số liệu của các tác giả khác, cơ quan tổ chức khác đều có trích dẫn "
            "và chú thích nguồn gốc.\n\n"
            "Nếu phát hiện có bất kỳ sự gian lận nào tôi xin hoàn toàn chịu trách nhiệm "
            "về nội dung đồ án của mình. "
            "Trường Đại học Tôn Đức Thắng không liên quan đến những vi phạm tác quyền, "
            "bản quyền do tôi gây ra trong quá trình thực hiện (nếu có)."
        )
        self.add_paragraph(text)

        self.doc.add_paragraph()
        sig = self.doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sig.paragraph_format.first_line_indent = None
        self._make_run(sig, f"TP. Hồ Chí Minh, ngày ... tháng ... năm {self.year}\n", italic=True)
        self._make_run(sig, "Tác giả\n(ký tên và ghi rõ họ tên)\n\n\n", italic=True)
        self._make_run(sig, self.author, italic=True)
        

    def add_abstract(self, vn_text, en_text):
        """Trang Tóm tắt (Tiếng Việt) + Abstract (Tiếng Anh)."""
        self.add_page_break()
        # Tóm tắt Tiếng Việt
        p_title = self.doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.first_line_indent = None
        self._make_run(p_title, self.title.upper(), bold=True, size=14)

        self.doc.add_paragraph()
        self.add_heading("TÓM TẮT")
        for para in vn_text.split('\n\n'):
            if para.strip():
                self.add_paragraph(para.strip())
        self.add_page_break()

        # Abstract Tiếng Anh
        self.add_heading("ABSTRACT")
        for para in en_text.split('\n\n'):
            if para.strip():
                self.add_paragraph(para.strip())
        

    # ═══════════════ HEADINGS ═══════════════
    def add_heading(self, text, level=1):
        """Heading chuẩn TDTU.
        level=1: Tên chương — 16pt, Bold, IN HOA
        level=2: Tiểu mục cấp 1 — 14pt, Bold
        level=3: Tiểu mục cấp 2 — 13pt, Bold
        """
        display_text = text.upper() if level == 1 else text
        
        # Luôn luôn ngắt trang trước Chương (Heading 1) theo đúng quy định TDTU
        if level == 1:
            self.add_page_break()
            
        h = self.doc.add_heading(display_text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h.paragraph_format.first_line_indent = None
        
        # Do đã ngắt trang thủ công ở trên hoặc không cần ngắt, ta set cái này = False
        # để tránh Word tự chèn thêm trang phụ do property page_break_before
        h.paragraph_format.page_break_before = False
        self._just_started_section = False

        size_map = {1: self.DEFAULTS["heading1_size"],
                    2: self.DEFAULTS["heading2_size"],
                    3: self.DEFAULTS["heading3_size"]}

        for run in h.runs:
            run.font.name = self.DEFAULTS["font_name"]
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.size = Pt(size_map.get(level, 13))
            run.bold = True

        # Track chương hiện tại
        if level == 1:
            match = re.search(r'CHƯƠNG\s+(\d+)', text, re.IGNORECASE)
            if match:
                self.current_chapter = int(match.group(1))

    # ═══════════════ PARAGRAPHS ═══════════════
    def add_paragraph(self, text, bold=False, italic=False, alignment=None):
        """Thêm đoạn văn bản chuẩn + fix spacing. Tránh lỗi canh lề Justify làm giãn chữ (Word soft-return)."""
        paragraphs_created = []
        for line in text.split('\n'):
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = self.DEFAULTS["font_name"]
            run.font.size = Pt(self.DEFAULTS["font_size"])
            run.bold = bold
            run.italic = italic
            self._fix_char_spacing(run)
            if alignment:
                p.alignment = alignment
            paragraphs_created.append(p)
        return paragraphs_created[-1] if paragraphs_created else None

    def add_component_table(self, components, chapter_num=3):
        """Bảng liệt kê linh kiện kèm placeholder ảnh từng cái.
        components: list of dict {name, description, specs}
        """
        for i, comp in enumerate(components, 1):
            self.add_heading(f"{chapter_num}.2.{i} {comp['name']}", level=3)
            self.add_paragraph(comp.get('description', ''))
            if comp.get('specs'):
                data = [['Thông số', 'Giá trị']]
                for k, v in comp['specs'].items():
                    data.append([k, v])
                self.add_table(data, caption=f"Bảng {chapter_num}.{i}: Thông số kỹ thuật {comp['name']}")
            self.add_placeholder_image(
                caption=f"Hình {chapter_num}.{i}: Linh kiện {comp['name']} sử dụng trong đề tài",
                width_inches=3.5
            )

    def add_bullet_list(self, items, bullet_char=""):
        """Thêm danh sách (Không dùng gạch ngang '-' theo yêu cầu)."""
        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.left_indent = Cm(1.27)
            prefix = f"{bullet_char} " if bullet_char else ""
            self._make_run(p, f"{prefix}{item}")

    # ═══════════════ FORMULAS ═══════════════
    def add_formula(self, formula_text, formula_number=None, explanation=None):
        """Công thức toán học. Tự động render bằng CodeCogs API. (Đã bỏ cột đánh số theo yêu cầu)."""
        import urllib.parse
        import requests
        import io

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        
        # Render công thức bằng ảnh
        cleaned_formula = formula_text.strip()
        success_img = False
        
        # Tạo URL an toàn, dpi150 và \Large cho kích thước đồng đều đẹp mắt
        safe_tex = urllib.parse.quote(cleaned_formula)
        url = f"https://latex.codecogs.com/png.image?\\dpi{{150}}\\bg{{white}}\\Large {safe_tex}"
        
        try:
            resp = requests.get(url, timeout=10)
            if resp.status_code == 200 and len(resp.content) > 100:
                img_stream = io.BytesIO(resp.content)
                try:
                    from PIL import Image
                    with Image.open(img_stream) as img:
                        w_px, h_px = img.size
                    img_stream.seek(0)
                    
                    # codecogs \dpi{150} maps approx to 150 pixels per inch.
                    expected_w_inches = w_px / 150.0
                    if expected_w_inches > 5.5:
                        expected_w_inches = 5.5 # Max width to prevent cutoff
                        
                    run = p.add_run()
                    run.add_picture(img_stream, width=Inches(expected_w_inches))
                    success_img = True
                except Exception:
                    img_stream.seek(0)
                    run = p.add_run()
                    run.add_picture(img_stream)
                    success_img = True
        except:
            pass
            
        if not success_img: # Fallback text
            self._make_run(p, formula_text, italic=True)
            
        self.doc.add_paragraph() # Spacing
        
        # Giải thích
        if explanation:
            self.add_paragraph(explanation)

    # ═══════════════ LẤY ẢNH TỪ WEB ═══════════════
    def add_image_from_url(self, url, caption=None, width_inches=None):
        """Tải và chèn ảnh trực tiếp từ đường link mxh hoặc web (VD: Logo TDTU)."""
        try:
            resp = requests.get(url, timeout=15)
            if resp.status_code == 200:
                img_stream = io.BytesIO(resp.content)
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(6)
                
                run = p.add_run()
                if width_inches:
                    run.add_picture(img_stream, width=Inches(width_inches))
                else:
                    run.add_picture(img_stream)
                    
                if caption:
                    cap_p = self.doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    self._make_run(cap_p, caption, italic=True)
        except Exception as e:
            print(f"[Warning] Failed to fetch layout image {url}: {e}")

    # ═══════════════ TABLES ═══════════════
    def add_table(self, data, caption=None, header=True):
        """Bảng biểu chuẩn TDTU — tiêu đề phía TRÊN."""
        if caption:
            cap_p = self.doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.keep_with_next = True # FIX: Keep caption stick to table
            cap_p.paragraph_format.first_line_indent = None
            self._make_run(cap_p, caption, bold=True)

        if not data:
            return

        rows, cols = len(data), len(data[0])
        table = self.doc.add_table(rows=rows, cols=cols, style='Table Grid')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Tắt tính năng break trang giữa chừng cho các dòng của bảng
        for row in table.rows:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trPr.append(parse_xml(r'<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

        for i, row_data in enumerate(data):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_text)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.first_line_indent = None
                    if i < rows - 1: # Giữ nguyên vẹn toàn bộ bảng trên 1 trang 
                        para.paragraph_format.keep_with_next = True
                    for run in para.runs:
                        run.font.name = self.DEFAULTS["font_name"]
                        run.font.size = Pt(self.DEFAULTS["font_size"])
                        if header and i == 0:
                            run.bold = True

        self.doc.add_paragraph()  # spacing

    # ═══════════════ CODE SNIPPETS ═══════════════
    def add_code(self, code_text, language="", caption=None):
        """Chèn code snippet với font Consolas 10pt + nền xám nhạt."""
        if caption:
            cap_p = self.doc.add_paragraph()
            cap_p.paragraph_format.first_line_indent = None
            self._make_run(cap_p, caption, bold=True, italic=True)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.5)

        # Nền xám nhạt
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
        p._element.get_or_add_pPr().append(shading)

        run = p.add_run(code_text)
        run.font.name = self.DEFAULTS["code_font"]
        run.font.size = Pt(self.DEFAULTS["code_size"])

        self.doc.add_paragraph()

    # ═══════════════ DIAGRAMS ═══════════════
    def add_mermaid_diagram(self, mermaid_code, caption="So do", width_inches=6.0):
        """Render Mermaid -> PNG -> chen vao Word. URL-safe encoding + retry + fallback.
        QUAN TRONG: mermaid_code PHAI chi dung ky tu ASCII/English. KHONG dung tieng Viet."""
        # Clean mermaid code: strip leading whitespace from each line
        lines = mermaid_code.strip().split('\n')
        cleaned = '\n'.join(line.strip() for line in lines)

        success = False
        last_err = ""

        for attempt in range(3):
            try:
                raw_b64 = base64.b64encode(cleaned.encode('utf-8')).decode('ascii')
                # URL-safe encode base64 string (fix +, /, = chars causing 404)
                safe_b64 = urllib.parse.quote(raw_b64, safe='')
                url = f"https://mermaid.ink/img/{safe_b64}?width=1400&bgColor=white"

                resp = requests.get(url, timeout=25)
                if resp.status_code == 200 and len(resp.content) > 500:
                    img_stream = io.BytesIO(resp.content)
                    self.doc.add_picture(img_stream, width=Inches(width_inches))
                    last_p = self.doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    last_p.paragraph_format.keep_with_next = True # Ngăn gãy trang chia lìa ảnh và caption

                    cap_p = self.doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.paragraph_format.first_line_indent = None
                    self._make_run(cap_p, caption, italic=True)
                    success = True
                    break
                else:
                    last_err = f"HTTP {resp.status_code}"
            except Exception as e:
                last_err = str(e)

            if attempt < 2:
                time.sleep(2)  # Wait before retry

        if not success:
            # Fallback: tao khung placeholder thay vi de loi
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0"/>')
            p._element.get_or_add_pPr().append(shading)
            run = self._make_run(
                p,
                f"\n\n[ SO DO SE DUOC VE BANG TAY - Xem file baocao.md ]\n(Loi: {last_err})\n\n",
                bold=True
            )
            run.font.color.rgb = RGBColor(120, 120, 120)

            cap_p = self.doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.first_line_indent = None
            self._make_run(cap_p, caption, italic=True)
            safe_cap = caption.encode('ascii', 'ignore').decode('ascii')
            safe_err = str(last_err).encode('ascii', 'ignore').decode('ascii')
            print(f"  [WARN] Mermaid fallback for: {safe_cap} ({safe_err})")

    def add_image(self, image_path, caption="Hình", width_inches=5.0):
        """Chèn hình ảnh local + chú thích phía dưới."""
        if os.path.exists(image_path):
            self.doc.add_picture(image_path, width=Inches(width_inches))
            last_p = self.doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_p.paragraph_format.keep_with_next = True

            cap_p = self.doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.paragraph_format.first_line_indent = None
            self._make_run(cap_p, caption, italic=True)
        else:
            self.add_paragraph(f"[Không tìm thấy hình: {image_path}]")

    def add_placeholder_image(self, caption="Hình", height_inches=3.0, width_inches=5.0):
        """Khung chứa ảnh trống để tác giả tự chèn ảnh vào sau."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = None
        
        # Tạo khung chứa nền xám
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E0E0E0"/>')
        p._element.get_or_add_pPr().append(shading)
        p.paragraph_format.keep_with_next = True
        
        run = self._make_run(p, f"\n\n[ KHUNG CHÈN ẢNH VỊ TRÍ NÀY ]\n(Nhấn vào đây để dán ảnh linh kiện/sản phẩm thực tế)\n\n\n", bold=True)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        cap_p = self.doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.first_line_indent = None
        self._make_run(cap_p, caption, italic=True)
        self.add_empty_lines(1)

    def generate_table_of_contents(self):
        """Thêm mục lục tự động (TOC). Không in dong huong dan ma ep Word tu dong f9 khi mo len."""
        self.add_page_break()
        self.add_page_break()
        self.add_heading("MỤC LỤC")
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = parse_xml(rf'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        instrText = parse_xml(rf'<w:instrText {nsdecls("w")} xml:space="preserve">TOC \o "1-3" \h \z \u</w:instrText>')
        fldChar2 = parse_xml(rf'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        fldChar3 = parse_xml(rf'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        
        # Ép file Word tự động update field khi người dùng mở file lên
        try:
            settings_element = self.doc.settings.element
            update_fields = parse_xml(r'<w:updateFields xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="true"/>')
            settings_element.append(update_fields)
        except Exception:
            pass
            
        

    # ═══════════════ REFERENCES ═══════════════
    def add_references(self, vn_refs=None, en_refs=None):
        """Thêm danh mục tài liệu tham khảo (APA6, chia TV/TA)."""
        self.add_page_break()
        
        self.add_heading("TÀI LIỆU THAM KHẢO")

        if vn_refs:
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = None
            self._make_run(p, "Tiếng Việt", bold=True)
            for ref in vn_refs:
                rp = self.doc.add_paragraph()
                rp.paragraph_format.first_line_indent = None
                rp.paragraph_format.left_indent = Cm(1.0)
                # Hanging indent
                self._make_run(rp, ref)

        if en_refs:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = None
            self._make_run(p, "Tiếng Anh", bold=True)
            for ref in en_refs:
                rp = self.doc.add_paragraph()
                rp.paragraph_format.first_line_indent = None
                rp.paragraph_format.left_indent = Cm(1.0)
                self._make_run(rp, ref)

    # ═══════════════ ABBREVIATION TABLE ═══════════════
    def add_abbreviation_table(self, abbreviations):
        """Thêm Danh mục chữ viết tắt (xếp A-Z).
        
        abbreviations: dict {"AMR": "Autonomous Mobile Robot", ...}
        """
        self.add_heading("DANH MỤC CÁC CHỮ VIẾT TẮT")
        sorted_abbrs = sorted(abbreviations.items())
        data = [["Viết tắt", "Ý nghĩa"]]
        for abbr, meaning in sorted_abbrs:
            data.append([abbr, meaning])
        self.add_table(data)

    # ═══════════════ PROJECT STRUCTURE ═══════════════
    def add_project_structure(self, root_dir, max_depth=3, max_files=15):
        """Tự động quét và chèn cấu trúc thư mục dự án."""
        self.add_heading("Cấu trúc thư mục dự án", level=2)
        structure = ""
        for root, dirs, files in os.walk(root_dir):
            level = root.replace(root_dir, '').count(os.sep)
            if level > max_depth:
                continue
            indent = '│   ' * level
            folder = os.path.basename(root)
            structure += f"{indent}├── {folder}/\n"
            subindent = '│   ' * (level + 1)
            for f in files[:max_files]:
                structure += f"{subindent}├── {f}\n"
            if len(files) > max_files:
                structure += f"{subindent}└── ... ({len(files) - max_files} files more)\n"

        self.add_code(structure, caption="Cấu trúc thư mục mã nguồn:")

    # ═══════════════ UTILITIES ═══════════════
    def add_page_break(self):
        """Ngắt trang (có kiểm soát để tránh trang trắng)."""
        if getattr(self, '_just_started_section', False):
            return
        if not self.doc.paragraphs:
            return
        last_p = self.doc.paragraphs[-1]
        if last_p.text.strip() == '' and 'w:br' in last_p._p.xml and 'w:type="page"' in last_p._p.xml:
            return
        self.doc.add_page_break()
    def add_empty_lines(self, count=1):
        """Thêm dòng trống."""
        for _ in range(count):
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = None

    def save(self, filename):
        """Lưu file .docx."""
        self.doc.save(filename)
        abs_path = os.path.abspath(filename)
        print(f"Bao cao da duoc luu tai: {abs_path}")
        print(f"   Kich thuoc: {os.path.getsize(abs_path) / 1024:.1f} KB")
        return abs_path


# ══════════════════════════════════════════════════════
# VÍ DỤ SỬ DỤNG (cho mọi đề tài)
# ══════════════════════════════════════════════════════
if __name__ == "__main__":
    gen = ReportGenerator(
        title="[TÊN ĐỀ TÀI CỦA BẠN]",
        author="[HỌ VÀ TÊN SINH VIÊN]",
        advisor="[HỌC HÀM. HỌ TÊN GVHD]",
        department="KHOA ĐIỆN – ĐIỆN TỬ",
        major="KỸ THUẬT ĐIỆN TỬ - VIỄN THÔNG",
        year="2026"
    )

    # === TRANG BÌA ===
    gen.add_cover_page(include_advisor=False)  # Bìa chính
    gen.add_cover_page(include_advisor=True)   # Bìa phụ

    # === FRONT MATTER ===
    gen.add_acknowledgment()
    gen.add_declaration()
    gen.add_abstract(
        vn_text="Đồ án này trình bày việc thiết kế và xây dựng...",
        en_text="This thesis presents the design and implementation of..."
    )
    gen.add_abbreviation_table({
        "IoT": "Internet of Things",
        "PID": "Proportional-Integral-Derivative",
        "MCU": "Microcontroller Unit",
    })

    # === NỘI DUNG ===
    gen.add_heading("CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI")
    gen.add_heading("1.1 Đặt vấn đề", level=2)
    gen.add_paragraph(
        "Trong bối cảnh cuộc cách mạng công nghiệp lần thứ tư, "
        "việc ứng dụng các công nghệ tiên tiến vào sản xuất và đời sống "
        "trở thành xu hướng tất yếu. Đề tài này tập trung vào..."
    )

    # Bảng mẫu (tiêu đề phía trên)
    gen.add_table(
        [["Thông số", "Giá trị", "Đơn vị"],
         ["Điện áp cấp", "12", "V"],
         ["Dòng tiêu thụ", "2", "A"],
         ["Công suất", "24", "W"]],
        caption="Bảng 1.1: Thông số kỹ thuật hệ thống"
    )

    # Công thức mẫu
    gen.add_formula(
        "v = (R / 2) × (vR + vL)",
        formula_number="3.1",
        explanation=(
            "Trong đó:\n"
            "- v: vận tốc tịnh tiến của robot (m/s)\n"
            "- R: bán kính bánh xe (m)\n"
            "- vR, vL: vận tốc góc bánh phải và bánh trái (rad/s)"
        )
    )

    # Code mẫu
    gen.add_code(
        'void setup() {\n  Serial.begin(115200);\n  pinMode(LED, OUTPUT);\n}',
        language="cpp",
        caption="Đoạn mã 4.1: Khởi tạo phần cứng"
    )

    # Sơ đồ Mermaid
    gen.add_mermaid_diagram(
        "graph TD; A[Cảm biến] --> B[Vi điều khiển]; B --> C[Cơ cấu chấp hành]; B --> D[Màn hình]",
        caption="Hình 4.1. Sơ đồ khối hệ thống"
    )

    # Tài liệu tham khảo
    gen.add_references(
        vn_refs=[
            'Nguyễn, V. A. (2020). Giáo trình Vi xử lý và Vi điều khiển. Hà Nội: NXB ĐHQG.',
        ],
        en_refs=[
            'Smith, R. (2010). Rethinking teacher education. Sydney: AACLM Press.',
            'Dempsey, I. (2012). Individual education programs. Australasian J. of Special Education, 36(1), 21-31.',
        ]
    )

    gen.save("BaoCao_Mau_TDTU.docx")
